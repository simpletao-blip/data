"""Build journal-style DOCX manuscripts from the evidence-locked Markdown files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)

MAIN_SECTIONS = [
    "abstract_draft.md",
    "keywords_draft.md",
    "introduction_draft.md",
    "methods_draft.md",
    "results_draft.md",
    "discussion_draft.md",
    "conclusion_draft.md",
    "code_data_availability.md",
    "declarations_draft.md",
    "references_draft.md",
]

AUTHOR_NAME = "Jiachao Tang (唐家超)"
AFFILIATION = "College of Automotive Engineering, Jilin University, Changchun 130022, China"
CORRESPONDING_EMAIL = "tangjc2224@mails.jlu.edu.cn"
ORCID = "0009-0000-5472-5902"

FIGURES = [
    (1, "figure1_workflow.png", "figure1_contract.md"),
    (2, "figure2_data_coverage.png", "figure2_contract.md"),
    (3, "figure3_mechanism_errors.png", "figure3_contract.md"),
    (4, "figure4_ensemble_validation.png", "figure4_contract.md"),
    (5, "figure5_reactivity_maps.png", "figure5_contract.md"),
    (6, "figure6_emissions_pareto.png", "figure6_contract.md"),
    (7, "figure7_sensitivity.png", "figure7_contract.md"),
]

SUPPLEMENTARY_FIGURES = [
    (1, "figureS1_literature_evidence.png"),
    (2, "figureS2_data_coverage.png"),
    (3, "figureS3_idt_criterion.png"),
    (4, "figureS4_mechanism_sizes.png"),
    (5, "figureS5_lbv_campaign_residuals.png"),
    (6, "figureS6_idt_residual_structure.png"),
    (7, "figureS7_grid_independence.png"),
    (8, "figureS8_surrogate_calibration.png"),
    (9, "figureS9_psr_fraction_sensitivity.png"),
    (10, "figureS10_heat_loss_sensitivity.png"),
    (11, "figureS11_pareto_support.png"),
    (12, "figureS12_reactor_paths.png"),
    (13, "figureS13_revision_robustness.png"),
]


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, value, end):
        run._r.append(element)


def configure_document(doc: Document, running_label: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(running_label)
    set_font(run, size=9, color=GRAY)
    add_page_field(section.footer.paragraphs[0])


def add_title_page(doc: Document, title: str, subtitle: str):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    set_font(r, size=24, bold=True, color=RGBColor(24, 45, 66))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(subtitle)
    set_font(r, size=13, italic=True, color=GRAY)
    for label, value in (
        ("Author", AUTHOR_NAME),
        ("Affiliation", AFFILIATION),
        ("Corresponding author", f"Jiachao Tang; {CORRESPONDING_EMAIL}"),
        ("ORCID", ORCID),
        ("Target journal", "Fuel"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        set_font(r, size=10.5, bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5)
    doc.add_page_break()


def add_inline_markup(paragraph, text: str, size=11):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Consolas", size=max(9, size - 1))
        elif token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size)


def editable_equation_text(source: str) -> str:
    """Convert the manuscript's display equations to editable Word text."""
    normalized = " ".join(source.split())
    equations = {
        r"(1-\alpha)\mathrm{NH_3}+1.5\alpha\mathrm{H_2}+0.5\alpha\mathrm{N_2},":
            "(1 − α)NH₃ + 1.5αH₂ + 0.5αN₂,",
        r"\tau_{\mathrm{ign}}=\operatorname*{arg\,max}_t\left(\frac{\mathrm{d}T}{\mathrm{d}t}\right),":
            "τ_ign = arg max_t(dT/dt),",
        r"E_{\mathrm{IDT}}=\left|\log_{10}\left(\tau_{\mathrm{sim}}/\tau_{\mathrm{exp}}\right)\right|.":
            "E_IDT = |log₁₀(τ_sim/τ_exp)|.",
        r"E_{\mathrm{LBV}}=\frac{|S_{u,\mathrm{sim}}-S_{u,\mathrm{exp}}|}{S_{u,\mathrm{exp}}}.":
            "E_LBV = |S_u,sim − S_u,exp| / S_u,exp.",
        r"s_j=\frac{\partial\ln S_u}{\partial\ln k_j}.":
            "s_j = ∂ln S_u / ∂ln k_j.",
        r"\dot m=m_{\mathrm{PSR}}/\tau_{\mathrm{PSR}}.":
            "ṁ = m_PSR / τ_PSR.",
        r"EI_k=10^9\frac{Y_{k,\mathrm{out}}\dot m}{\dot Q_{\mathrm{fuel}}}\quad\mathrm{g\,MJ^{-1}},":
            "EI_k = 10⁹(Y_k,out ṁ / Q̇_fuel)  g MJ⁻¹,",
        r"EI_{\mathrm{NOx}}=\frac{M_{\mathrm{NO_2}}}{M_{\mathrm{NO}}}EI_{\mathrm{NO}}+EI_{\mathrm{NO_2}}.":
            "EI_NOx = (M_NO₂/M_NO)EI_NO + EI_NO₂.",
        r"\min_{\mathbf w}\frac{1}{N}\sum_n\left(\frac{\sum_i w_i p_{ni}-y_n}{q_n}\right)^2, \quad w_i\geq0,\quad\sum_iw_i=1.":
            "min_w (1/N) Σ_n [( Σ_i w_i p_ni − y_n )/q_n]²,   w_i ≥ 0,   Σ_i w_i = 1.",
        r"\Delta H_{\mathrm{crack}}=\alpha\left(1.5h_{\mathrm{H_2}}+0.5h_{\mathrm{N_2}}-h_{\mathrm{NH_3}}\right).":
            "ΔH_crack = α(1.5h_H₂ + 0.5h_N₂ − h_NH₃).",
    }
    return equations.get(normalized, normalized)


def add_inline_main_figure(doc: Document, number: int):
    _, image_name, contract_name = next(item for item in FIGURES if item[0] == number)
    caption_text = contract_caption(number, ROOT / "figures" / contract_name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    add_picture_with_alt(
        p.add_run(),
        ROOT / "figures" / "exports" / image_name,
        Inches(6.35),
        caption_text,
    )
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run(caption_text)
    set_font(r, size=9.5)


def add_markdown(doc: Document, text: str, inserted_figures: set[int] | None = None):
    lines = text.splitlines()
    paragraph_buffer: list[str] = []
    equation_buffer: list[str] = []
    in_equation = False

    def flush_paragraph():
        if paragraph_buffer:
            paragraph_text = " ".join(line.strip() for line in paragraph_buffer)
            p = doc.add_paragraph()
            add_inline_markup(p, paragraph_text)
            paragraph_buffer.clear()
            if inserted_figures is not None:
                for number, _, _ in FIGURES:
                    if number in inserted_figures:
                        continue
                    if re.search(rf"\bFig(?:ure)?\.?\s*{number}(?!\d|[A-Za-z])", paragraph_text):
                        add_inline_main_figure(doc, number)
                        inserted_figures.add(number)

    for raw in lines:
        line = raw.rstrip()
        if line.strip() == r"\[":
            flush_paragraph()
            in_equation = True
            equation_buffer = []
            continue
        if line.strip() == r"\]" and in_equation:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_font(
                p.add_run(editable_equation_text(" ".join(equation_buffer))),
                name="Cambria Math",
                size=10.5,
            )
            in_equation = False
            equation_buffer = []
            continue
        if in_equation:
            equation_buffer.append(line.strip())
            continue
        if not line.strip():
            flush_paragraph()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text_value = re.sub(r"\s+[—-]\s+(working|evidence-locked|reproducible).*draft$", "", heading.group(2), flags=re.I)
            doc.add_heading(text_value, level=level)
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.194)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline_markup(p, bullet.group(1))
            continue
        reference = re.match(r"^(\d+)\.\s+(.+)$", line)
        if reference:
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.35)
            p.paragraph_format.space_after = Pt(4)
            add_inline_markup(p, f"{reference.group(1)}. {reference.group(2)}", size=9.5)
            continue
        paragraph_buffer.append(line)
    flush_paragraph()


def contract_caption(number: int, path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    parts = []
    for line in text.splitlines():
        line = line.strip().lstrip("-").strip()
        if line.lower().startswith("core conclusion:"):
            parts.append(line.split(":", 1)[1].strip())
        elif line.lower().startswith("evidence logic:") or line.lower().startswith("evidence chain:"):
            parts.append(line.split(":", 1)[1].strip())
    return f"Figure {number}. " + " ".join(parts)


def add_picture_with_alt(run, image_path: Path, width, alt_text: str):
    shape = run.add_picture(str(image_path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text.split(".", 1)[0])
    return shape


def add_main_figures(doc: Document):
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for index, (number, image_name, contract_name) in enumerate(FIGURES):
        if index:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        caption_text = contract_caption(number, ROOT / "figures" / contract_name)
        add_picture_with_alt(
            p.add_run(),
            ROOT / "figures" / "exports" / image_name,
            Inches(6.45),
            caption_text,
        )
        caption = doc.add_paragraph()
        caption.paragraph_format.space_before = Pt(6)
        caption.paragraph_format.space_after = Pt(0)
        caption.paragraph_format.keep_with_next = True
        r = caption.add_run(caption_text)
        set_font(r, size=9.5)


def build_main():
    doc = Document()
    configure_document(doc, "Partially cracked ammonia | Fuel manuscript")
    doc.core_properties.author = "Jiachao Tang"
    doc.core_properties.subject = "Kinetic-model uncertainty in partially cracked ammonia combustion"
    title = "Study-held-out validation exposes evidence gaps in multi-mechanism screening of partially cracked ammonia"
    add_title_page(doc, title, "Original research article")
    inserted_figures: set[int] = set()
    for filename in MAIN_SECTIONS:
        add_markdown(
            doc,
            (ROOT / "manuscript" / filename).read_text(encoding="utf-8"),
            inserted_figures,
        )
    missing = sorted(set(range(1, 8)) - inserted_figures)
    if missing:
        raise RuntimeError(f"Main figures were not cited and embedded: {missing}")
    path = OUT / "manuscript_fuel_submission.docx"
    doc.save(path)
    print(path)


def build_si():
    doc = Document()
    configure_document(doc, "Partially cracked ammonia | Supplementary Information")
    doc.core_properties.author = "Jiachao Tang"
    add_title_page(
        doc,
        "Supplementary Information",
        "Study-held-out validation exposes evidence gaps in multi-mechanism screening of partially cracked ammonia",
    )
    add_markdown(
        doc, (ROOT / "manuscript" / "supplementary_information_draft.md").read_text(encoding="utf-8")
    )
    for number, image_name in SUPPLEMENTARY_FIGURES:
        doc.add_page_break()
        doc.add_heading(f"Supplementary Figure S{number}", level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_text = contract_caption(
            number, ROOT / "figures" / f"figureS{number}_contract.md"
        ).replace(f"Figure {number}", f"Figure S{number}")
        add_picture_with_alt(
            p.add_run(),
            ROOT / "figures" / "supplementary" / image_name,
            Inches(6.45),
            caption_text,
        )
        caption = doc.add_paragraph()
        add_inline_markup(caption, caption_text, size=9.5)
    path = OUT / "supplementary_information_fuel_submission.docx"
    doc.save(path)
    print(path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_main()
    build_si()


if __name__ == "__main__":
    main()
